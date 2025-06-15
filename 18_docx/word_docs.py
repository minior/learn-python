#https://python-docx.readthedocs.io/en/latest/

import docx

# Document obj -> contains Paragraph objs -> contains Run objs

d = docx.Document('18_docx\\my_resume_copy.docx') # returns document object
print(d.paragraphs) # visualise contained para objects
print()
print(d.paragraphs[1].text) # my heading
p = d.paragraphs[30] # saving index 30 para object (uni of leeds entry)

print()
print(p.runs) # visualise contained run objects (when style changes eg. Bold)
print()
print(p.runs[3].text) # wonky, but this is where it is BOLDED
print(p.runs[3].bold)
print(p.runs[3].italic)

# we can edit the doc this way
p.runs[3].underline = True
p.runs[3].text = "ABCABCABC"
d.paragraphs[30].style = 'Title'
d.save('18_resume_edit.docx')

# create a blank docx object
d = docx.Document()
d.add_paragraph("Created Paragraph 1.")
d.add_paragraph("Another One")

# add to existing para
d.paragraphs[0].add_run(" This is a created Run")
d.paragraphs[0].runs[1].bold = True
d.save('18_docx\\create_new.docx')

# limitation: can't insert text between indexes -> copy existing & insert between text copies